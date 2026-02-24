"""
ingestion/parser.py — Parse any supported document type to text + metadata.

Supported: .pdf  .docx  .txt  .md  .html  .csv
"""
from __future__ import annotations

import csv
import io
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

import structlog

logger = structlog.get_logger(__name__)


@dataclass
class ParsedDocument:
    text: str
    pages: List[str] = field(default_factory=list)
    tables: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


def parse(file_path: str) -> ParsedDocument:
    """Dispatch to the correct parser based on file extension."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = path.suffix.lower()
    parsers = {
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".txt": _parse_text,
        ".md": _parse_text,
        ".html": _parse_html,
        ".htm": _parse_html,
        ".csv": _parse_csv,
    }

    parser_fn = parsers.get(ext, _parse_text)
    logger.info("parsing_file", file=str(path), ext=ext)

    try:
        result = parser_fn(str(path))
        result.metadata["file_name"] = path.name
        result.metadata["file_ext"] = ext
        result.metadata["file_size"] = path.stat().st_size
        return result
    except Exception as exc:
        logger.error("parse_failed", file=str(path), error=str(exc))
        raise


# ── PDF ───────────────────────────────────────────────────────────────────────

# Minimum characters to consider a page "has real text" (not noise)
_OCR_TEXT_THRESHOLD = 20

# Lazy singleton — easyocr reader loads model only once across all pages/calls
_easyocr_reader = None


def _get_ocr_reader():
    global _easyocr_reader
    if _easyocr_reader is None:
        try:
            import easyocr  # type: ignore
            logger.info("ocr_model_loading", engine="easyocr")
            _easyocr_reader = easyocr.Reader(["en"], gpu=False, verbose=False)
            logger.info("ocr_model_ready")
        except ImportError:
            logger.warning("easyocr_not_installed", msg="pip install easyocr")
            _easyocr_reader = False  # sentinel: don't retry
    return _easyocr_reader if _easyocr_reader is not False else None


def _ocr_page(page) -> str:
    """
    Render a PyMuPDF page to a numpy array and run easyocr.
    Called only for image-only / scanned pages.
    """
    reader = _get_ocr_reader()
    if reader is None:
        return ""

    import numpy as np  # type: ignore

    # Render at 150 DPI — good quality/speed balance
    mat = page.get_pixmap(dpi=150)
    img = np.frombuffer(mat.samples, dtype=np.uint8).reshape(mat.height, mat.width, mat.n)
    if mat.n == 4:
        img = img[:, :, :3]

    result = reader.readtext(img, detail=0)  # detail=0 → list of strings, faster
    return "\n".join(result)


def _parse_pdf(file_path: str) -> ParsedDocument:
    """
    Fast PDF parser using PyMuPDF (fitz).
    - Native text PDFs: ~5-50ms/page
    - Scanned/image PDFs: automatic OCR fallback via easyocr (model loaded once)
    """
    try:
        import fitz  # type: ignore  (PyMuPDF)
    except ImportError:
        raise ImportError("PyMuPDF is required for PDF parsing: pip install pymupdf")

    pages: List[str] = []
    tables: List[str] = []
    ocr_pages = 0

    doc = fitz.open(file_path)
    try:
        # First pass: identify which pages need OCR, render all at once if needed
        page_texts: List[str] = []
        scanned_indices: List[int] = []

        for i, page in enumerate(doc):
            txt = page.get_text("text") or ""
            page_texts.append(txt)
            if len(txt.strip()) < _OCR_TEXT_THRESHOLD:
                scanned_indices.append(i)

        if scanned_indices:
            logger.info(
                "pdf_scanned_detected",
                scanned_pages=len(scanned_indices),
                total_pages=doc.page_count,
            )
            # Pre-warm the OCR model before looping pages (loads weights once)
            _get_ocr_reader()

        for i, page in enumerate(doc):
            if i in scanned_indices:
                page_text = _ocr_page(page)
                ocr_pages += 1
            else:
                page_text = page_texts[i]
                # Lightweight table detection for native-text pages
                for block in page.get_text("blocks"):
                    if block[6] != 0:  # skip image blocks
                        continue
                    blk_text = block[4].strip()
                    lines = blk_text.splitlines()
                    if len(lines) >= 2 and sum(len(l.split()) >= 3 for l in lines) >= 2:
                        tables.append(blk_text)

            pages.append(page_text)
    finally:
        doc.close()

    if ocr_pages:
        logger.info("pdf_ocr_complete", ocr_pages=ocr_pages, total_pages=len(pages))

    full_text = "\n\n".join(pages)
    return ParsedDocument(
        text=full_text,
        pages=pages,
        tables=tables,
        metadata={"page_count": len(pages), "ocr_pages": ocr_pages},
    )



# ── DOCX ──────────────────────────────────────────────────────────────────────

def _parse_docx(file_path: str) -> ParsedDocument:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables: List[str] = []

    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        tables.append("\n".join(rows))

    full_text = "\n\n".join(paragraphs)
    if tables:
        full_text += "\n\n" + "\n\n[TABLE]\n".join(tables)

    return ParsedDocument(
        text=full_text,
        pages=[full_text],
        tables=tables,
        metadata={"paragraph_count": len(paragraphs)},
    )


# ── Plain text / Markdown ─────────────────────────────────────────────────────

def _parse_text(file_path: str) -> ParsedDocument:
    with open(file_path, encoding="utf-8", errors="replace") as fh:
        text = fh.read()
    return ParsedDocument(text=text, pages=[text], metadata={})


# ── HTML ──────────────────────────────────────────────────────────────────────

def _parse_html(file_path: str) -> ParsedDocument:
    try:
        from bs4 import BeautifulSoup  # type: ignore
    except ImportError:
        raise ImportError("beautifulsoup4 is required: pip install beautifulsoup4")

    with open(file_path, encoding="utf-8", errors="replace") as fh:
        raw = fh.read()

    soup = BeautifulSoup(raw, "html.parser")
    # Remove scripts and style tags
    for tag in soup(["script", "style"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    # Collapse excessive blank lines
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    text = "\n".join(lines)

    return ParsedDocument(text=text, pages=[text], metadata={})


# ── CSV ───────────────────────────────────────────────────────────────────────

def _parse_csv(file_path: str) -> ParsedDocument:
    rows: List[List[str]] = []
    with open(file_path, newline="", encoding="utf-8", errors="replace") as fh:
        reader = csv.reader(fh)
        for row in reader:
            rows.append(row)

    if not rows:
        return ParsedDocument(text="", metadata={})

    header = rows[0]
    table_lines = [" | ".join(header)]
    for row in rows[1:]:
        table_lines.append(" | ".join(row))
    table_text = "\n".join(table_lines)

    return ParsedDocument(
        text=table_text,
        pages=[table_text],
        tables=[table_text],
        metadata={"row_count": len(rows) - 1, "column_count": len(header)},
    )
